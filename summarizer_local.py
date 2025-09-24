import ollama
from docx import Document

# Read job listings from a .docx file
def read_job_listings(file_path):
    """
    Reads the content from a .docx file and returns it as a single string.
    """
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def create_output_docx(summary, output_path):
    doc = Document()
    doc.add_heading('Job Listings Summary', 0)
    doc.add_paragraph(summary)
    doc.save(output_path)

def call_ollama_api(prompt):
    response = ollama.chat(model='llama3.2:1b', messages=[
        {
            'role': 'user',
            'content': prompt,
        },
    ])
    return response['message']['content']

# Summarize the text using Ollama and save the summary to a new .docx file
def summarize_and_save(input_file, output_file):
    """
    Summarizes job listings using the local Ollama LLM and saves the summary.
    """
    try:
        # Read the job listings from the input file
        job_listings = read_job_listings(input_file)
        
        if not job_listings.strip():
            print("The input file is empty. Nothing to summarize.")
            return

        # Define the prompt for the LLM
        prompt = f"""
        Summarize the following job listings in a concise manner.
        Make three lists : 
        1. Of all skills
        2. Of expected skills
        3. Of nice-to-have skills 
        Mentioned more than once. Count each category and rank them separately. 
        Then provide the summary in a clear and structured format.
        
        Job Listings:
        {job_listings}
        """

        print("Sending job listings to Ollama for summarization...")
        
        summary = call_ollama_api(prompt)
        print("Received summary from Ollama.")
        
        # Create a new Word document and save the summary
        create_output_docx(summary, output_file)
        print(f"Successfully summarized and saved the result to '{output_file}'.")

    except FileNotFoundError:
        print(f"Error: The file '{input_file}' was not found.")
    except Exception as e:
        print(f"An error occurred: {e}")

# Main execution block
if __name__ == "__main__":
    input_file_path = "ofertyPracy.docx"
    output_file_path = "summary.docx"
    
    # Run the function to summarize and save
    summarize_and_save(input_file_path, output_file_path)